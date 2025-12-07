import json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from typing import List, Dict, Any
from datetime import datetime

# Load criteria definitions
tev1_szov = ["0-150 hallgató","200-300 hallgató","350-500 hallgató","550-700 hallgató","750+ hallgató"]
tev1_pont = [3, 6, 10, 14, 15]
tev2_szov =["A tevékenység vonatkozásában megfogalmazott célok nem vagy elhanyagolható mértékben tükrözték az elvárt eredmény(eke)t", "A tevékenység vonatkozásában megfogalmazott célok átlagos mértékben tükrözték az elvárt eredmény(eke)t" ,"A tevékenység vonatkozásában megfogalmazott célok átlagon felüli mértékben tükrözték az elvárt eredmény(eke)t", "A tevékenység vonatkozásában megfogalmazott célok kiválóan / teljes mértékben tükrözték az elvárt eredmény(eke)t"] 
tev2_pont =[2, 6, 8, 10]
tev3_szov = ["A tevékenység megvalósítása nem vagy elhanyagolható mértékben igényel speciális tudást vagy háttérismeretet és a tevékenység nem komplex","A tevékenység megvalósítása átlagos mértékű speciális tudást igényel és komplex a tevékenység","A tevékenység megvalósítása jelentős mértékű háttérismeretet, speciális tudást igényel és komplex a tevékenység"]
tev3_pont = [2, 6, 10]
tev4_szov = ["A határidőket nem, vagy csak többszöri felszólításra tartotta be", "A határidőket egyszeri felszólításra tartotta be", "A határidőket önmagától betartotta"]
tev4_pont = [1, 3, 5]
tev5_szov = ["A tevékenységet önállóan nem vagy alig tudta megvalósítani", "A tevékenységet részleges önállósággal tudta megvalósítani", "A tevékenységet teljes önállósággal tudta megvalósítani"]
tev5_pont = [1, 3, 5]
tev6_szov = ["A tevékenységet végző személy nem vagy elhanyagolható mértékben fordított időt a tevékenység megvalósítására", "A tevékenységet végző személy átlagos mértékben fordított időt a tevékenység megvalósítására", "A tevékenységet végző személy  jelentős időmennyiséget fordított a tevékenység megvalósítására"]
tev6_pont = [1, 6, 8]
tev7_szov = ["Jelen szempont a közéleti tevékenység végzése szempontjából nem releváns","az esetek többségében inaktív viselkedést tanúsított.", "az esetek többségében átlagosan aktív viselkedést tanúsított.", "az esetek többségében proaktív viselkedést tanúsított."]
tev7_pont = [0, 2, 6, 8]

CRITERIA = [
    {"name": "1. Hallgatók száma", "szov": tev1_szov, "pont": tev1_pont},
    {"name": "2. Célok megvalósulása", "szov": tev2_szov, "pont": tev2_pont},
    {"name": "3. Szükséges tudás", "szov": tev3_szov, "pont": tev3_pont},
    {"name": "4. Határidő betartása", "szov": tev4_szov, "pont": tev4_pont},
    {"name": "5. Önállóság", "szov": tev5_szov, "pont": tev5_pont},
    {"name": "6. Fordított idő", "szov": tev6_szov, "pont": tev6_pont},
    {"name": "7. Viselkedés", "szov": tev7_szov, "pont": tev7_pont},
]


def get_grade_ranges(points_list: List[int]) -> List[tuple]:
    """Convert points list to grade ranges."""
    if not points_list:
        return []
    
    ranges = []
    prev = 0
    for p in points_list:
        if p < prev:
            continue
        ranges.append((prev, p))
        prev = p + 1
    
    return ranges


def point_to_grade_index(point_value: int, points_list: List[int]) -> int:
    """Convert actual point value to grade index (0-based)."""
    if point_value is None or point_value == 0:
        return -1
    
    ranges = get_grade_ranges(points_list)
    for idx, (min_p, max_p) in enumerate(ranges):
        if min_p <= point_value <= max_p:
            return idx
    
    return -1


def get_grade_text(point_value: int, criteria_dict: Dict) -> str:
    """Get the grade description for a given point value."""
    if point_value is None or point_value == 0:
        return criteria_dict['szov'][0]
    
    pts = criteria_dict['pont']
    texts = criteria_dict['szov']
    grade_idx = point_to_grade_index(point_value, pts)
    
    if grade_idx >= 0 and grade_idx < len(texts):
        return texts[grade_idx]
    else:
        return texts[0] if texts else "N/A"


def load_scores(scores_file: Path = Path('scores.json')) -> List[Dict[str, Any]]:
    """Load scores from JSON file."""
    if not scores_file.exists():
        return []
    
    try:
        with open(scores_file, 'r', encoding='utf-8') as f:
            return json.load(f)
    except Exception as e:
        print(f"Error loading scores: {e}")
        return []


def fill_template(student: Dict[str, Any], template_path: Path, output_path: Path):
    """Fill Word template with student data."""
    try:
        # Load template
        doc = Document(template_path)
        
        # Replace placeholders in document
        # Student info
        replace_text_in_document(doc, '[NÉV]', student['name'])
        replace_text_in_document(doc, '[NEPTUN]', student['neptun'])
        replace_text_in_document(doc, '[SZAK]', student['major'])
        replace_text_in_document(doc, '[ÉV]', student['time'])
        
        # Calculate total
        total = 0
        for i, p in enumerate(student['pontszam']):
            if p is not None and student['is_relevant'][i]:
                total += p
        
        replace_text_in_document(doc, '[ÖSSZ_PONT]', str(total))
        
        # Max total
        max_total = sum(max(c['pont']) for c in CRITERIA)
        replace_text_in_document(doc, '[MAX_PONT]', str(max_total))
        
        # Description/notes
        leiras = student.get('leiras', '')
        replace_text_in_document(doc, '[LEIRAS]', leiras)
        
        # Criteria scores
        for i, crit in enumerate(CRITERIA):
            point_value = student['pontszam'][i]
            is_relevant = student['is_relevant'][i]
            
            # Replace point value
            replace_text_in_document(doc, f'[TEV{i+1}_PONT]', str(point_value) if point_value is not None else '0')
            
            # Replace grade text
            if is_relevant and point_value is not None:
                grade_text = get_grade_text(point_value, crit)
            else:
                grade_text = "Nem releváns" if not is_relevant else "Nincs választás"
            
            replace_text_in_document(doc, f'[TEV{i+1}_SZOV]', grade_text)
        
        # Save filled document
        doc.save(output_path)
        print(f"Document saved: {output_path}")
        return True
        
    except Exception as e:
        print(f"Error filling template: {e}")
        return False


def replace_text_in_document(doc: Document, old_text: str, new_text: str):
    """Replace all occurrences of text in document (including in tables, paragraphs, etc.)."""
    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        if old_text in paragraph.text:
            replace_text_in_paragraph(paragraph, old_text, new_text)
    
    # Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if old_text in paragraph.text:
                        replace_text_in_paragraph(paragraph, old_text, new_text)


def replace_text_in_paragraph(paragraph, old_text: str, new_text: str):
    """Replace text within a single paragraph while preserving formatting."""
    if old_text not in paragraph.text:
        return
    
    # Clear the paragraph
    for run in paragraph.runs:
        run.text = ""
    
    # Add new text
    paragraph.text = paragraph.text.replace(old_text, new_text)


def generate_filled_documents(template_path: Path = Path('sablon.docx'), 
                               scores_file: Path = Path('scores.json'),
                               output_dir: Path = Path('filled_documents')):
    """Generate filled documents for all students."""
    
    # Create output directory
    output_dir.mkdir(exist_ok=True)
    
    # Load scores
    students = load_scores(scores_file)
    
    if not students:
        print("No students found in scores.json")
        return
    
    # Check template exists
    if not template_path.exists():
        print(f"Template not found: {template_path}")
        return
    
    # Process each student
    for student in students:
        output_file = output_dir / f"{student['name']}_{student['neptun']}.docx"
        print(f"Processing {student['name']}...")
        fill_template(student, template_path, output_file)
    
    print(f"\nAll documents generated in: {output_dir}")


if __name__ == '__main__':
    # Generate filled documents for all students
    generate_filled_documents()
