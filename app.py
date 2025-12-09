import streamlit as st
import pandas as pd
from typing import List, Dict, Any
import json
import os
import subprocess
from pathlib import Path

# --- Criteria definitions (copied from main.py) ---
tev1_szov = ["0-150 hallgató","200-300 hallgató","350-500 hallgató","550-700 hallgató","750+ hallgató"]
tev1_pont = [3, 6, 10, 14, 15]
tev2_szov =["A tevékenység vonatkozásában megfogalmazott célok nem vagy elhanyagolható mértékben tükrözték az elvárt eredmény(eke)t", "A tevékenység vonatkozásában megfogalmazott célok átlagos mértékben tükrözték az elvárt eredmény(eke)t" ,"A tevékenység vonatkozásában megfogalmazott célok átlagon felüli mértékben tükrözték az elvárt eredmény(eke)t", "A tevékenység vonatkozásában megfogalmazott célok kiválóan / teljes mértékben tükrözték az elvárt eredmény(eke)t"] 
tev2_pont =[2, 6, 8, 10]
tev3_szov = ["A tevékenység megvalósítása nem vagy elhanyagolható mértékben igényel speciális tudást vagy háttérismeretet és a tevékenység nem komplex","A tevékenység megvalósítása átlagos mértékű speciális tudást igényel és komplex a tevékenység","A tevékenység megvalósítása jelentős mértékű háttérismeretet, speciális tudást igényel és komplex a tevékenység"]
tev3_pont = [2, 6, 10]
tev4_szov = ["A határidőket nem, vagy csak többszöri felszólításra tartotta be", "A határidőket egyszeri felszólításra tartotta be", "A határidőket önmagától betartotta"]
tev4_pont = [1, 3, 5]
tev5_szov = ["A tevékenységet önállóan nem vagy alig tudta megvalósítani", "A tevékenységet részleges önállósággal tudta megvalósítani", "A tevékenységet teljes önállósággal tudta megvalósítani"]
tev5_pont = [1, 3, 5]
tev6_szov = ["A tevékenységet végző személy nem vagy elhanyagolható mértékben fordított időt a tevékenység megvalósítására", "A tevékenységet végző személy átlagos mértékben fordított időt a tevékenység megvalósítására", "A tevékenységet végző személy  jelentős időmennyiséget fordított a tevékenység megvalósítására"]
tev6_pont = [1, 6, 8]
tev7_szov = ["Jelen szempont a közéleti tevékenység végzése szempontjából nem releváns","az esetek többségében inaktív viselkedést tanúsított.", "az esetek többségében átlagosan aktív viselkedést tanúsított.", "az esetek többségében proaktív viselkedést tanúsított."]
tev7_pont = [0, 2, 6, 8]

CRITERIA = [
    {"name": "1. Hallgatók száma", "szov": tev1_szov, "pont": tev1_pont},
    {"name": "2. Célok megvalósulása", "szov": tev2_szov, "pont": tev2_pont},
    {"name": "3. Szükséges tudás", "szov": tev3_szov, "pont": tev3_pont},
    {"name": "4. Határidő betartása", "szov": tev4_szov, "pont": tev4_pont},
    {"name": "5. Önállóság", "szov": tev5_szov, "pont": tev5_pont},
    {"name": "6. Fordított idő", "szov": tev6_szov, "pont": tev6_pont},
    {"name": "7. Viselkedés", "szov": tev7_szov, "pont": tev7_pont},
]

# Persistence file path
SCORES_FILE = Path('scores.json')


def save_scores(students: List[Dict[str, Any]]):
    """Save student scores to JSON file."""
    data = []
    for s in students:
        data.append({
            'name': s['name'],
            'time': s['time'],
            'neptun': s['neptun'],
            'major': s['major'],
            'pontszam': s['pontszam'],
            'is_relevant': s['is_relevant'],
            'is_done': s['is_done'],
            'leiras': s.get('leiras', ''),
            'kategoria': s.get('kategoria', ''),
        })
    with open(SCORES_FILE, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def load_scores(students: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Load saved scores from JSON file and merge with current student list."""
    if not SCORES_FILE.exists():
        return students
    
    try:
        with open(SCORES_FILE, 'r', encoding='utf-8') as f:
            saved_data = json.load(f)
        
        # Create a map of neptun -> saved scores
        saved_map = {item['neptun']: item for item in saved_data}
        
        # Merge saved data into current students
        for s in students:
            if s['neptun'] in saved_map:
                saved = saved_map[s['neptun']]
                s['pontszam'] = saved.get('pontszam', [None] * len(CRITERIA))
                s['is_relevant'] = saved.get('is_relevant', [True] * len(CRITERIA))
                s['is_done'] = saved.get('is_done', False)
                s['leiras'] = saved.get('leiras', '')
                s['kategoria'] = saved.get('kategoria', '')
        
        return students
    except Exception as e:
        st.warning(f"Hiba a scores.json betöltésekor: {e}")
        return students


def get_grade_ranges(points_list: List[int]) -> List[tuple]:
    """
    Convert points list to grade ranges.
    E.g., [3, 6, 10, 14, 15] → [(0, 3), (4, 6), (7, 10), (11, 14), (15, 15)]
    """
    if not points_list:
        return []
    
    ranges = []
    prev = 0
    for p in points_list:
        if p < prev:
            continue  # Skip if out of order
        ranges.append((prev, p))
        prev = p + 1
    
    return ranges


def point_to_grade_index(point_value: int, points_list: List[int]) -> int:
    """Convert actual point value to grade index (0-based)."""
    if point_value is None or point_value == 0:
        return -1  # No selection
    
    ranges = get_grade_ranges(points_list)
    for idx, (min_p, max_p) in enumerate(ranges):
        if min_p <= point_value <= max_p:
            return idx
    
    return -1


def grade_index_to_point(grade_idx: int, points_list: List[int]) -> int:
    """Convert grade index to the corresponding maximum point value."""
    if grade_idx < 0 or grade_idx >= len(points_list):
        return 0
    return points_list[grade_idx]



def load_students(path: str) -> List[Dict[str, Any]]:
    df = pd.read_csv(path, sep=';')
    students = []
    for _, row in df.iterrows():
        students.append({
            'name': row.get('name', ''),
            'time': row.get('time', ''),
            'neptun': row.get('neptun', ''),
            'major': row.get('major', ''),
            'pontszam': [None] * len(CRITERIA),
            'is_relevant': [True] * len(CRITERIA),
            'is_done': False,
            'leiras': '',
            'kategoria': '',
        })
    return students


def calculate_total(st: Dict[str, Any]) -> int:
    total = 0
    for i, p in enumerate(st['pontszam']):
        if p is not None and st['is_relevant'][i]:
            total += p
    return total


def calculate_osszeg(total_points: int) -> int:
    """Calculate OSSZEG (total money): points × 2500"""
    return total_points * 2500


def ensure_state(path='data.csv'):
    if 'students' not in st.session_state:
        students = load_students(path)
        # Load any previously saved scores
        students = load_scores(students)
        st.session_state['students'] = students


def student_dataframe(students: List[Dict[str, Any]]) -> pd.DataFrame:
    rows = []
    for s in students:
        row = {
            'name': s['name'],
            'time': s['time'],
            'neptun': s['neptun'],
            'major': s['major'],
            'is_done': s['is_done'],
        }
        for i in range(len(CRITERIA)):
            row[f'p_{i}'] = s['pontszam'][i]
            row[f'rel_{i}'] = s['is_relevant'][i]
        rows.append(row)
    return pd.DataFrame(rows)


def get_grade_ranges(points_list: List[int]) -> List[tuple]:
    """Convert points list to grade ranges for filling Word doc."""
    if not points_list:
        return []
    
    ranges = []
    prev = 0
    for p in points_list:
        if p < prev:
            continue
        ranges.append((prev, p))
        prev = p + 1
    
    return ranges


def point_to_grade_index_word(point_value: int, points_list: List[int]) -> int:
    """Convert actual point value to grade index for Word filling."""
    if point_value is None or point_value == 0:
        return -1
    
    ranges = get_grade_ranges(points_list)
    for idx, (min_p, max_p) in enumerate(ranges):
        if min_p <= point_value <= max_p:
            return idx
    
    return -1


def get_grade_text_word(point_value: int, criteria_dict: Dict) -> str:
    """Get grade text for Word document."""
    if point_value is None or point_value == 0:
        return criteria_dict['szov'][0]
    
    pts = criteria_dict['pont']
    texts = criteria_dict['szov']
    grade_idx = point_to_grade_index_word(point_value, pts)
    
    if grade_idx >= 0 and grade_idx < len(texts):
        return texts[grade_idx]
    else:
        return texts[0] if texts else "N/A"


def fill_word_document(student: Dict[str, Any], template_path: Path = Path('sablon.docx')):
    """Fill Word template with student data and return path to filled document.

    This collects all inserted values and then calls a helper that transfers
    paragraph formatting from the preceding blank paragraph before deleting it.
    That preserves the template's visual style while removing the extra blank rows.
    """
    from docx import Document
    import zipfile
    import tempfile
    import shutil

    try:
        output_dir = Path('filled_documents')
        output_dir.mkdir(exist_ok=True)

        # First, do direct XML replacement on the template
        temp_dir = Path(tempfile.mkdtemp())
        temp_template = temp_dir / 'template.docx'
        shutil.copy(template_path, temp_template)
        
        # Build replacement map
        replacements = {}
        replacements['[NÉV]'] = student['name']
        replacements['[NEPTUN]'] = student['neptun']
        replacements['[SZAK]'] = student['major']
        replacements['[ÉV]'] = student['time']
        replacements['[KATEGORIA]'] = student.get('kategoria', '')
        
        total = 0
        for i, p in enumerate(student['pontszam']):
            if p is not None and student['is_relevant'][i]:
                total += p
        
        replacements['[ÖSSZ_PONT]'] = str(total)
        max_total = sum(max(c['pont']) for c in CRITERIA)
        replacements['[MAX_PONT]'] = str(max_total)
        osszeg = calculate_osszeg(total)
        replacements['[OSSZEG]'] = str(osszeg)
        replacements['[LEIRAS]'] = student.get('leiras', '')
        
        for i, crit in enumerate(CRITERIA):
            point_value = student['pontszam'][i]
            is_relevant = student['is_relevant'][i]
            
            if not is_relevant:
                pt_str = '0'
                grade_text = "Jelen szempont a közéleti tevékenység végzése szempontjából nem releváns"
            elif point_value is not None:
                pt_str = str(point_value)
                grade_text = get_grade_text_word(point_value, crit)
            else:
                pt_str = '0'
                grade_text = "Nincs választás"
            
            replacements[f'[TEV{i+1}_PONT]'] = pt_str
            replacements[f'[TEV{i+1}_SZOV]'] = grade_text
        
        # Extract docx, replace in XML files, repack
        with zipfile.ZipFile(temp_template, 'r') as zip_ref:
            zip_ref.extractall(temp_dir / 'docx_contents')
        
        # Replace in all XML files
        found_placeholders = set()
        for xml_file in (temp_dir / 'docx_contents').rglob('*.xml'):
            try:
                content = xml_file.read_text(encoding='utf-8')
                # Check which placeholders are actually in the file
                for placeholder in replacements.keys():
                    if placeholder in content:
                        found_placeholders.add(placeholder)
                
                for placeholder, value in replacements.items():
                    content = content.replace(placeholder, str(value))
                xml_file.write_text(content, encoding='utf-8')
            except Exception:
                pass
        
        # Debug: show which placeholders were found
        not_found = set(replacements.keys()) - found_placeholders
        if not_found:
            st.warning(f"⚠️ Ezek a helyőrzők NEM találhatók a sablonban: {', '.join(not_found)}")
            st.info(f"✓ Ezek a helyőrzők MEGTALÁLHATÓK: {', '.join(found_placeholders) if found_placeholders else 'EGYIK SEM'}")
            st.info("💡 Kérlek ellenőrizd, hogy a Word sablonban pontosan ezek a karakterláncok szerepelnek-e (pl. [NEPTUN], [ÉV], stb.)")
        
        # Repack as docx
        output_file = output_dir / f"{student['name']}_{student['neptun']}.docx"
        with zipfile.ZipFile(output_file, 'w', zipfile.ZIP_DEFLATED) as docx:
            for file in (temp_dir / 'docx_contents').rglob('*'):
                if file.is_file():
                    arcname = file.relative_to(temp_dir / 'docx_contents')
                    docx.write(file, arcname)
        
        # Cleanup temp directory
        shutil.rmtree(temp_dir)
        
        return output_file

    except Exception as e:
        st.error(f"Hiba a Word dokumentum kitöltésekor: {e}")
        return None


def replace_text_in_doc(doc, old_text: str, new_text: str):
    """Replace placeholders without losing run-level styling."""

    def replace_in_paragraph(paragraph):
        replaced = False
        # Fast path when the placeholder sits wholly inside single runs
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
                replaced = True
        if replaced:
            return

        # Fallback for placeholders split across multiple runs
        full_text = "".join(run.text for run in paragraph.runs)
        if old_text not in full_text:
            return

        start = full_text.find(old_text)
        end = start + len(old_text)

        current = 0
        inserted = False
        parts_by_run = {run: [] for run in paragraph.runs}

        for run in paragraph.runs:
            text = run.text
            length = len(text)
            run_end = current + length

            if run_end <= start or current >= end:
                parts_by_run[run].append(text)
            else:
                prefix_len = max(0, start - current)
                suffix_len = max(0, run_end - end)
                if prefix_len:
                    parts_by_run[run].append(text[:prefix_len])
                if not inserted:
                    parts_by_run[run].append(new_text)
                    inserted = True
                if suffix_len:
                    parts_by_run[run].append(text[length - suffix_len:])
            current += length

        for run in paragraph.runs:
            run.text = "".join(parts_by_run.get(run, []))

    def walk_paragraphs(paragraphs):
        for paragraph in paragraphs:
            if old_text in paragraph.text:
                try:
                    replace_in_paragraph(paragraph)
                except Exception:
                    pass

    def walk_tables(tables):
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    walk_paragraphs(cell.paragraphs)
                    walk_tables(cell.tables)

    # Body
    walk_paragraphs(doc.paragraphs)
    walk_tables(doc.tables)

    # Headers and footers
    for section in doc.sections:
        walk_paragraphs(section.header.paragraphs)
        walk_tables(section.header.tables)
        walk_paragraphs(section.footer.paragraphs)
        walk_tables(section.footer.tables)

    # Text boxes (shapes) in headers/footers/sections where placeholders often live
    def walk_shapes(shapes):
        for shape in shapes:
            # Some shapes can contain text frames; if so, process their paragraphs/tables
            if hasattr(shape, "text_frame") and shape.text_frame is not None:
                walk_paragraphs(shape.text_frame.paragraphs)
                walk_tables(shape.text_frame.tables)

    for section in doc.sections:
        try:
            if hasattr(section, 'header') and hasattr(section.header, 'shapes'):
                walk_shapes(section.header.shapes)
        except Exception:
            pass
        try:
            if hasattr(section, 'footer') and hasattr(section.footer, 'shapes'):
                walk_shapes(section.footer.shapes)
        except Exception:
            pass

    # Fallback: replace in all text nodes (w:t) anywhere in the document tree
    try:
        namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        text_elements = doc.element.xpath('.//w:t', namespaces=namespaces)
        for t in text_elements:
            if t.text and old_text in t.text:
                t.text = t.text.replace(old_text, new_text)
    except Exception as e:
        # Silently continue if xpath fails
        pass
    
    # Additional fallback: try without namespace (for some Word versions)
    try:
        for element in doc.element.iter():
            if element.text and old_text in element.text:
                element.text = element.text.replace(old_text, new_text)
            if element.tail and old_text in element.tail:
                element.tail = element.tail.replace(old_text, new_text)
    except Exception:
        pass


def transfer_format_and_remove_prev_blank(doc, filled_values: List[str]):
    """No-op placeholder to avoid formatting-related errors during export."""
    return


def remove_preceding_newline(doc, placeholder: str):
    """Remove newline character before a placeholder. (Deprecated)"""
    pass


def main():
    st.set_page_config(layout='wide', page_title='Közéleti pontozó', initial_sidebar_state='expanded')
    
    # Compact CSS: reduce margins and padding
    # Theme toggle (top-left sidebar)
    with st.sidebar:
        top_left = st.container()
    with top_left:
        if 'theme' not in st.session_state:
            st.session_state['theme'] = 'dark'
        theme_choice = st.radio('Téma', ['🌙 Dark', '☀️ Light'], index=0 if st.session_state['theme']=='dark' else 1)
        st.session_state['theme'] = 'dark' if 'Dark' in theme_choice else 'light'

    if st.session_state['theme'] == 'dark':
        palette = {
            'bg': '#0f172a',
            'card': '#111827',
            'muted': '#94a3b8',
            'text': '#e2e8f0',
            'accent': '#22d3ee',
            'border': '#1f2937',
            'button_text': '#0b1220'
        }
    else:
        palette = {
            'bg': '#f8fafc',
            'card': '#ffffff',
            'muted': '#475569',
            'text': '#0f172a',
            'accent': '#0ea5e9',
            'border': '#e2e8f0',
            'button_text': '#ffffff'
        }

    css = f"""
    <style>
    :root {{
        --bg: {palette['bg']};
        --card: {palette['card']};
        --muted: {palette['muted']};
        --text: {palette['text']};
        --accent: {palette['accent']};
        --border: {palette['border']};
    }}
    body, .stApp, .main, [data-testid="stSidebar"], [data-testid="stHeader"], [data-testid="stToolbar"] {{
        background: var(--bg) !important;
        color: var(--text) !important;
    }}
    .block-container {{ padding-top: 1rem; }}
    .element-container {{ margin-bottom: 0.5rem !important; }}
    [data-testid="stVerticalBlock"] > [style*="flex-direction: column"] > [data-testid="stVerticalBlock"] {{
        gap: 0 !important;
    }}
    h1, h2, h3, h4, h5, h6, label, p, span, div {{
        color: var(--text) !important;
    }}
    .stButton>button {{
        background: var(--accent) !important;
        color: {palette['button_text']} !important;
        border: none !important;
    }}
    .stButton>button:hover {{ filter: brightness(0.95); }}
    .stCheckbox>label, .stRadio>div>label {{ color: var(--text) !important; }}
    .stTextInput>div>div>input, .stTextArea textarea, .stNumberInput input {{
        background: var(--card) !important;
        color: var(--text) !important;
        border: 1px solid var(--border) !important;
    }}
    .stMetric {{ background: var(--card) !important; border: 1px solid var(--border); padding: 0.5rem; border-radius: 0.5rem; }}
    .stMarkdown, .stCaption {{ color: var(--muted) !important; }}
    </style>
    """
    st.markdown(css, unsafe_allow_html=True)
    
    st.title('Közéleti tevékenység pontozó')

    ensure_state('data.csv')
    students = st.session_state['students']

    # Define helper callbacks early
    import re

    def make_slider_callback(sidx_local, idx_local, sel_key_local):
        def _cb():
            val = st.session_state.get(sel_key_local, 0)
            # treat 0 as no selection
            st.session_state['students'][sidx_local]['pontszam'][idx_local] = int(val) if int(val) != 0 else None
            # Save to disk immediately without rerun
            save_scores(st.session_state['students'])
        return _cb

    def make_rel_callback(sidx_local, idx_local, rel_key_local):
        def _cb():
            val = st.session_state.get(rel_key_local, True)
            st.session_state['students'][sidx_local]['is_relevant'][idx_local] = bool(val)
            save_scores(st.session_state['students'])
        return _cb

    def make_done_callback(sidx_local, done_key_local):
        def _cb():
            val = st.session_state.get(done_key_local, False)
            st.session_state['students'][sidx_local]['is_done'] = bool(val)
            save_scores(st.session_state['students'])
        return _cb
    
    def make_leiras_callback(sidx_local, desc_key_local):
        def _cb():
            val = st.session_state.get(desc_key_local, '')
            st.session_state['students'][sidx_local]['leiras'] = str(val)[:500]
            save_scores(st.session_state['students'])
        return _cb
    
    def make_kategoria_callback(sidx_local, kat_key_local):
        def _cb():
            val = st.session_state.get(kat_key_local, '')
            st.session_state['students'][sidx_local]['kategoria'] = str(val)[:5]
            save_scores(st.session_state['students'])
        return _cb

    def _grade_slider_callback(sidx_local, idx_local, sel_key_local, pts_local):
        """Callback for numeric input: store the entered point value."""
        point_val = st.session_state.get(sel_key_local, 0)
        if point_val is None or point_val < 0:
            st.session_state['students'][sidx_local]['pontszam'][idx_local] = None
        else:
            # Clamp to max valid point
            max_pt = max(pts_local) if pts_local else 0
            point_val = min(point_val, max_pt)
            st.session_state['students'][sidx_local]['pontszam'][idx_local] = int(point_val)
        
        save_scores(st.session_state['students'])
    with st.sidebar:
        st.header('Hallgatók')
        opts = [f"{s['name']} — {calculate_total(s)}p ({calculate_osszeg(calculate_total(s))}){' ✓' if s['is_done'] else ''}" for s in students]

        if opts:
            # Persist selected student index in session state even across reruns
            if 'selected_student_idx' not in st.session_state:
                st.session_state['selected_student_idx'] = 0

            sel = st.radio(
                'Válassz hallgatót',
                options=list(range(len(opts))),
                index=st.session_state['selected_student_idx'],
                format_func=lambda i: opts[i],
                key='sel_radio'
            )
            st.session_state['selected_student_idx'] = sel
        else:
            st.info('Nincs betöltött hallgató')
            return

        st.markdown(f"Összesen: **{len(students)}**")

    sidx = sel
    
    # Header row with student info and metrics inline
    header_cols = st.columns([3, 1, 1, 1, 1])
    with header_cols[0]:
        st.markdown(f"## {students[sidx]['name']}")
        st.caption(f"Neptun: {students[sidx]['neptun']} — Szak: {students[sidx]['major']} — Év: {students[sidx]['time']}")
    
    with header_cols[1]:
        total = calculate_total(students[sidx])
        st.metric('Összpontszám', f"{total}p")
    
    with header_cols[2]:
        max_total = sum(max(c['pont']) for c in CRITERIA)
        st.metric('Max', f"{max_total}p")
    
    with header_cols[3]:
        osszeg = calculate_osszeg(calculate_total(students[sidx]))
        st.metric('OSSZEG', f"{osszeg}")
    
    with header_cols[4]:
        done_key = f'done_{sidx}'
        if done_key not in st.session_state:
            st.session_state[done_key] = students[sidx]['is_done']
        st.checkbox('Kész', value=st.session_state[done_key], key=done_key, on_change=make_done_callback(sidx, done_key))

    st.markdown('---')

    # Criteria

    def get_description_for_value(crit_local, val_local):
        if val_local is None or val_local == 0:
            return 'Nincs választás'
        pts = crit_local['pont']
        texts = crit_local['szov']
        # find highest index where point <= val_local
        idx = 0
        for j, p in enumerate(pts):
            if val_local >= p:
                idx = j
        # clamp
        if idx < 0:
            idx = 0
        if idx >= len(texts):
            idx = len(texts) - 1
        return texts[idx]

    for i, crit in enumerate(CRITERIA):
        st.markdown(f"**{crit['name']}**")
        col1, col2 = st.columns([3, 1])
        
        rel_key = f'rel_{sidx}_{i}'
        sel_key = f'sel_{sidx}_{i}'
        
        # relevance checkbox (narrow column)
        with col2:
            if rel_key not in st.session_state:
                st.session_state[rel_key] = students[sidx]['is_relevant'][i]
            is_relevant = st.checkbox('Rel.', value=st.session_state[rel_key], key=rel_key, on_change=make_rel_callback(sidx, i, rel_key))
        
        # numeric input + description (wide column)
        with col1:
            if not is_relevant:
                # Show "not relevant" message
                st.caption("⚠️ Jelen szempont a közéleti tevékenység végzése szempontjából nem releváns")
            else:
                pts = crit['pont']
                texts = crit['szov']
                max_pt = max(pts) if pts else 0
                
                # Current point value stored
                current_point = students[sidx]['pontszam'][i]
                if current_point is None:
                    current_point = 0
                
                if sel_key not in st.session_state:
                    st.session_state[sel_key] = int(current_point)
                
                input_col, desc_col = st.columns([2, 3])
                with input_col:
                    # Numeric input from 0 to max_pt
                    point_val = st.number_input(
                        'Pont',
                        min_value=0,
                        max_value=int(max_pt),
                        value=int(st.session_state.get(sel_key, 0)),
                        step=1,
                        key=sel_key,
                        on_change=lambda sidx_l=sidx, idx_l=i, sel_k=sel_key, pts_l=pts: _grade_slider_callback(sidx_l, idx_l, sel_k, pts_l),
                        label_visibility='collapsed'
                    )
                
                with desc_col:
                    # Map point value to grade text
                    # When 0, show first grade text; otherwise find based on threshold
                    if point_val == 0:
                        desc = texts[0] if texts else 'Nincs választás'
                    else:
                        grade_idx = point_to_grade_index(int(point_val), pts)
                        if grade_idx >= 0 and grade_idx < len(texts):
                            desc = texts[grade_idx]
                        else:
                            desc = texts[0] if texts else 'Nincs választás'
                    st.caption(desc)

    st.markdown('---')
    
    # Category field (max 2 characters)
    kat_key = f'kategoria_{sidx}'
    if kat_key not in st.session_state:
        st.session_state[kat_key] = students[sidx].get('kategoria', '')
    st.text_input(
        'Kategória',
        value=st.session_state.get(kat_key, ''),
        max_chars=2,
        key=kat_key,
        on_change=make_kategoria_callback(sidx, kat_key),
        placeholder='Max 2 karakter'
    )
    
    # Description field (single for entire student) - hidden but stored
    desc_key = f'leiras_{sidx}'
    if desc_key not in st.session_state:
        st.session_state[desc_key] = students[sidx]['leiras']
    st.text_area(
        'Leírás',
        value=st.session_state.get(desc_key, ''),
        max_chars=500,
        key=desc_key,
        on_change=make_leiras_callback(sidx, desc_key),
        placeholder='A Pályázó...',
        height=80,
        label_visibility='collapsed'
    )
    
    st.markdown('---')
    # Actions
    c1, c2 = st.columns(2)
    with c1:
        if st.button('Word doksi kitöltés'):
            filled_doc = fill_word_document(students[sidx])
            if filled_doc:
                st.success(f"Dokumentum elkészült: {filled_doc.name}")
                # Open the document
                try:
                    subprocess.Popen(['open', str(filled_doc)])
                except Exception as e:
                    st.warning(f"Hiba a dokumentum megnyitásakor: {e}")
    with c2:
        if st.button('Frissít listát'):
            st.success('Lista frissítve')


if __name__ == '__main__':
    main()

